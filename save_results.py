import docx

def save_blog_text(blog_text):
    document = docx.Document()
    document.add_heading('Блог о раннем обучении')
    document.add_paragraph(blog_text)
    document.save('blog_text.docx')

def save_image_text(image_text):
    with open('image_text.txt', 'w') as f:
        f.write(image_text)

# All modules save (incorrect) in "__pycache__"